"""
文件解析服务 - 支持 txt, pdf, docx, xlsx
"""
import logging
import re

logger = logging.getLogger(__name__)


class FileParser:
    """文件解析器"""
    
    SUPPORTED_EXTENSIONS = {'.txt', '.pdf', '.docx', '.xlsx', '.xls', '.md'}
    
    @classmethod
    def get_supported_extensions(cls) -> set:
        return cls.SUPPORTED_EXTENSIONS
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析TXT/MD文件"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'utf-16', 'latin-1']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                return content.strip()
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                logger.error(f"读取文件失败 {file_path}: {e}")
                raise
        
        raise ValueError(f"无法解析文件编码: {file_path}")
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析PDF文件 - 使用pymupdf，对中文支持更好"""
        try:
            import fitz  # pymupdf
            
            content_parts = []
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # 获取文本块并按位置排序
                blocks = page.get_text("dict")["blocks"]
                page_text = []
                
                for block in blocks:
                    if block["type"] == 0:  # 文本块
                        for line in block.get("lines", []):
                            line_text = ""
                            for span in line.get("spans", []):
                                line_text += span.get("text", "")
                            if line_text.strip():
                                page_text.append(line_text.strip())
                
                if page_text:
                    content_parts.append("\n".join(page_text))
            
            doc.close()
            content = "\n\n".join(content_parts)
            
            # 清理多余空白
            content = re.sub(r'\n{3,}', '\n\n', content)
            return content.strip()
            
        except ImportError:
            logger.warning("pymupdf未安装，使用备用方案")
            return FileParser._parse_pdf_fallback(file_path)
        except Exception as e:
            logger.warning(f"pymupdf解析失败: {e}，使用备用方案")
            return FileParser._parse_pdf_fallback(file_path)
    
    @staticmethod
    def _parse_pdf_fallback(file_path: str) -> str:
        """PDF备用解析方案"""
        content_parts = []
        
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
        except Exception as e:
            logger.warning(f"pdfplumber解析失败: {e}")
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text = page.extract_text()
                        if text:
                            content_parts.append(text)
            except Exception as e2:
                logger.error(f"所有PDF解析方案都失败 {file_path}: {e2}")
                raise
        
        return "\n\n".join(content_parts).strip()
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析Word文档"""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content_parts = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content_parts.append(text)
            
            # 也提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            return "\n".join(content_parts).strip()
            
        except ImportError:
            raise ValueError("python-docx未安装，无法解析Word文档")
        except Exception as e:
            logger.error(f"解析Word文档失败 {file_path}: {e}")
            raise
    
    @staticmethod
    def parse_xlsx(file_path: str) -> str:
        """解析Excel文件"""
        try:
            from openpyxl import load_workbook
            
            wb = load_workbook(file_path, read_only=True, data_only=True)
            content_parts = []
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_content = [f"[工作表: {sheet_name}]"]
                
                for row in sheet.iter_rows(values_only=True):
                    row_values = [str(cell) if cell is not None else "" for cell in row]
                    row_text = " | ".join(v for v in row_values if v)
                    if row_text:
                        sheet_content.append(row_text)
                
                if len(sheet_content) > 1:
                    content_parts.append("\n".join(sheet_content))
            
            wb.close()
            return "\n\n".join(content_parts).strip()
            
        except ImportError:
            raise ValueError("openpyxl未安装，无法解析Excel文件")
        except Exception as e:
            logger.error(f"解析Excel文件失败 {file_path}: {e}")
            raise
    
    @classmethod
    def parse(cls, file_path: str) -> str:
        """根据文件类型自动选择解析器"""
        import os
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext in ('.txt', '.md'):
            return cls.parse_txt(file_path)
        elif ext == '.pdf':
            return cls.parse_pdf(file_path)
        elif ext == '.docx':
            return cls.parse_docx(file_path)
        elif ext in ('.xlsx', '.xls'):
            return cls.parse_xlsx(file_path)
        else:
            raise ValueError(f"不支持的文件类型: {ext}")
